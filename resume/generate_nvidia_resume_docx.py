from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "Kyungtack_Lee_Resume_NVIDIA.docx"

NAVY = "17324D"
GREEN = "5A8F00"
DARK = "20262E"
MUTED = "56616D"
LIGHT = "D8DEE5"

# compact_reference_guide with a named resume override:
# Letter portrait; 0.58 in page margins; Arial 9.5 pt body; 1.05 line spacing;
# no layout tables; real Word numbering for bullets; quiet footer with page field.
PAGE_MARGIN = Inches(0.58)
CONTENT_WIDTH = Inches(8.5 - 1.16)


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_free_font(run, name="Arial", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = rgb(color)
    if italic is not None:
        run.italic = italic


def set_style_font(style, name, size, color=DARK, bold=False):
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold


def set_keep_with_next(paragraph, enabled=True):
    paragraph.paragraph_format.keep_with_next = enabled


def add_bottom_border(paragraph, color=GREEN, size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_cell_free_font(run, size=8, color=MUTED)


def add_hyperlink(paragraph, text, url, color=MUTED):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "17")
    r_pr.append(size)
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    r_pr.append(color_element)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "270")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "252")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, "Arial", 9.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color, bold, before, after in [
        ("Resume Name", 22, NAVY, False, 0, 0),
        ("Resume Tagline", 10.5, GREEN, True, 0, 2),
        ("Resume Contact", 8.5, MUTED, False, 0, 5),
        ("Resume Section", 11.5, NAVY, True, 7, 3),
        ("Resume Entry", 10, DARK, True, 3, 1),
        ("Resume Small", 8.7, DARK, False, 0, 1),
    ]:
        if name in styles:
            style = styles[name]
        else:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, "Arial", size, color=color, bold=bold)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.03
        style.paragraph_format.keep_with_next = name in {"Resume Section", "Resume Entry"}


def add_section_heading(doc, title):
    paragraph = doc.add_paragraph(style="Resume Section")
    paragraph.add_run(title.upper())
    add_bottom_border(paragraph)
    return paragraph


def add_entry_heading(doc, title, organization, dates):
    paragraph = doc.add_paragraph(style="Resume Entry")
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        CONTENT_WIDTH, WD_TAB_ALIGNMENT.RIGHT
    )
    paragraph.add_run(f"{title} | {organization}")
    date_run = paragraph.add_run(f"\t{dates}")
    set_cell_free_font(date_run, size=8.8, bold=True, color=MUTED)
    return paragraph


def add_bullet(doc, num_id, text, bold_prefix=None, small=False):
    paragraph = doc.add_paragraph(style="Resume Small" if small else "Normal")
    apply_bullet(paragraph, num_id)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_cell_free_font(prefix, size=8.7 if small else 9.5, bold=True, color=DARK)
        run = paragraph.add_run(text[len(bold_prefix) :])
        set_cell_free_font(run, size=8.7 if small else 9.5, color=DARK)
    else:
        run = paragraph.add_run(text)
        set_cell_free_font(run, size=8.7 if small else 9.5, color=DARK)
    return paragraph


def add_plain(doc, text, style="Normal", bold_prefix=None):
    paragraph = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_cell_free_font(run, size=9.5, bold=True, color=DARK)
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(0)
    add_bottom_border(paragraph, color=LIGHT, size="3", space="2")
    run = paragraph.add_run("Kyungtack Lee | Resume | Page ")
    set_cell_free_font(run, size=8, color=MUTED)
    add_page_field(paragraph)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = PAGE_MARGIN
    section.bottom_margin = Inches(0.68)
    section.left_margin = PAGE_MARGIN
    section.right_margin = PAGE_MARGIN
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.28)
    configure_styles(doc)
    bullet_num_id = create_bullet_numbering(doc)
    add_footer(section)

    name = doc.add_paragraph(style="Resume Name")
    name.add_run("Kyungtack Lee")
    tagline = doc.add_paragraph(style="Resume Tagline")
    tagline.add_run("VEHICLE MOTION CONTROL & AUTONOMOUS DRIVING ENGINEER")
    contact = doc.add_paragraph(style="Resume Contact")
    contact.add_run("Seoul, Republic of Korea | +82 10-2632-3242 | ")
    add_hyperlink(contact, "kyungtacklee@gmail.com", "mailto:kyungtacklee@gmail.com")
    contact.add_run(" | ")
    add_hyperlink(contact, "kyungtacklee.github.io", "https://kyungtacklee.github.io/")
    contact.add_run(" | ")
    add_hyperlink(contact, "github.com/kyungtacklee", "https://github.com/kyungtacklee")

    add_section_heading(doc, "Professional Summary")
    add_plain(
        doc,
        "Vehicle motion control and autonomous driving engineer with more than ten years of experience in production automotive engineering. Technical lead and project owner for safety-critical planning and control functions, with end-to-end responsibility spanning function architecture, algorithm development, modeling and simulation, real-time implementation, vehicle integration, calibration, and validation. Expertise includes motion planning, trajectory generation and tracking, vehicle dynamics, integrated chassis control, state estimation, and optimization-based control.",
    )

    add_section_heading(doc, "Core Competencies")
    skills = [
        ("Planning & Control: ", "Motion planning, trajectory generation, path tracking, integrated chassis control, optimization-based control"),
        ("Vehicle Engineering: ", "Vehicle dynamics, state estimation, actuator coordination, calibration, in-vehicle testing, recorded-data analysis"),
        ("Programming: ", "C for production software analysis and development; MATLAB; Python"),
        ("Real-Time Development: ", "MATLAB/Simulink, dSPACE MicroAutoBox II, real-time model build and deployment, rapid control prototyping"),
        ("NVIDIA Platforms: ", "NVIDIA Jetson AGX; Isaac Sim"),
        ("Simulation & Process: ", "CarSim, CarMaker, AUTOSAR-aligned development, unit/integration testing; familiarity with ISO 26262, ISO 21448 SOTIF, and Automotive SPICE"),
    ]
    for prefix, value in skills:
        add_bullet(doc, bullet_num_id, prefix + value, bold_prefix=prefix)

    add_section_heading(doc, "Professional Experience")
    add_entry_heading(doc, "Senior Research Engineer", "HL Mando", "2015.08 - Present")
    experience = [
        ("Mobility Motion Control (2020 - Present): ", "Technical lead and project owner for planning and control functions including Integrated Chassis Control, Vehicle Stability Control Assist, Evasive Collision Avoidance, Smart Hitching Assist, Trailer Parking Assist, and Minimum Risk Maneuver."),
        (None, "Lead function definition and algorithm development through simulation, real-time deployment, system integration, vehicle tuning, scenario-based testing, and measurable vehicle-level evaluation."),
        (None, "Develop and analyze production software in C; build and deploy MATLAB/Simulink control models to dSPACE MicroAutoBox II; use NVIDIA Jetson AGX for vehicle evaluation."),
        ("System Design (2019 - 2020): ", "E-Corner Module system design and analysis."),
        ("Gear Design (2015 - 2019): ", "Electric Power Steering, Steer-by-Wire, Shift-by-Wire, and e-Drive systems; durability analysis, optimization, and in-house engineering tool development."),
    ]
    for prefix, value in experience:
        add_bullet(doc, bullet_num_id, (prefix or "") + value, bold_prefix=prefix)
    add_entry_heading(
        doc,
        "Research Engineer",
        "Samsung Techwin (now Hanwha Aerospace)",
        "2012.08 - 2015.07",
    )
    add_bullet(
        doc,
        bullet_num_id,
        "Designed integral helical gears for turbo compressors and performed core-system design and dynamic-system analysis.",
    )
    add_entry_heading(doc, "Engineering Intern", "General Motors Korea", "2011.07 - 2011.08")
    add_bullet(doc, bullet_num_id, "Supported engine-map tuning and validation.")

    page_break = doc.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    add_section_heading(doc, "Selected Projects")
    projects = [
        (
            "Trailer Parking Assist & Minimum Risk Maneuver",
            "2026 - Present",
            [
                "Technical lead and project owner for planning and control architecture, simulation, real-time implementation, system integration, and vehicle-level evaluation.",
                "Developing robust maneuver generation and safe fallback behavior for production-oriented automated-driving functions.",
            ],
        ),
        (
            "Smart Hitching Assist",
            "2025",
            [
                "Led end-to-end function architecture, planning and control development, real-time deployment, vehicle integration, calibration, and validation.",
                "Delivered a successful customer demonstration; received a Company Special Recognition Award in December 2025.",
            ],
        ),
        (
            "Evasive Collision Avoidance",
            "2024",
            [
                "Owned the complete planning and control scope: evasive path generation, path tracking control, and vehicle stability coordination.",
                "Integrated and validated the unified framework through simulation, dSPACE real-time execution, in-vehicle tuning, and safety-critical scenario testing.",
            ],
        ),
        (
            "Vehicle Stability Control Assist",
            "2024",
            [
                "Developed a hierarchical integrated chassis controller combining differential braking and semi-active suspension damping.",
                "Published simulation and real-world evaluation results showing approximately 17.4% lower maximum roll angle and 8.7% lower maximum sideslip angle versus respective conventional methods.",
            ],
        ),
    ]
    for title, date, bullets in projects:
        add_entry_heading(doc, title, "HL Mando", date)
        for item in bullets:
            add_bullet(doc, bullet_num_id, item)

    add_section_heading(doc, "Education")
    add_entry_heading(
        doc,
        "Ph.D. Candidate, Mechanical Engineering",
        "Seoul National University",
        "2023.03 - Present",
    )
    add_plain(
        doc,
        "Interactive and Networked Robotics Laboratory (INRoL), Advisor: Professor Dongjun Lee",
        style="Resume Small",
    )
    add_plain(
        doc,
        "Research: Supervisory Integrated Chassis Control Using Model Predictive Path Integral Control with Legacy Controllers",
        style="Resume Small",
    )
    add_entry_heading(
        doc,
        "M.S., Mechanical Engineering",
        "Seoul National University",
        "2020.09 - 2022.06",
    )
    add_plain(
        doc,
        "Vehicle Dynamics and Control Laboratory (VDCL), Advisor: Professor Kyongsu Yi",
        style="Resume Small",
    )
    add_plain(
        doc,
        "Thesis: Path Tracking Control of Four-Wheel-Independent-Steering-and-Driving Vehicle Based on Adaptive-Weight Optimal Control",
        style="Resume Small",
    )
    add_entry_heading(doc, "B.S., Mechanical Engineering", "Ajou University", "2006.03 - 2012.06")

    add_section_heading(doc, "Selected Honors")
    honors = [
        "2025.12 | Company Special Recognition Award - Smart Hitching Assist development and customer demonstration, HL Mando",
        "2025.11 | Outstanding Paper Award (Oral Session), Korean Society of Automotive Engineers",
        "2024.12 | Excellence Award - R&D Outstanding Paper, HL Mando",
        "2024.11 | Outstanding Paper Award (Poster Session), Korean Society of Automotive Engineers",
        "2024.03 | Best Dialogue Award, EVS37 International Electric Vehicle Symposium and Exhibition",
    ]
    for item in honors:
        add_bullet(doc, bullet_num_id, item, small=True)

    add_section_heading(doc, "Selected Publications & Patents")
    publications = [
        'K. Lee and J. Seol, "Development of Integrated Chassis Control of Semi-Active Suspension with Differential Brake for Vehicle Lateral Stability," World Electric Vehicle Journal, 16(2):91, 2025.',
        'K. Lee et al., "Lyapunov-Informed Model Predictive Path Integral Control for Robust Trailer Hitch Assist under Perception Uncertainty," KSAE Annual Conference, 2025.',
        'K. Lee et al., "Continuous Curvature Path Planning Based on Bezier Curves for Autonomous Driving in Complex Environments," KSAE Annual Conference, 2024.',
        "Inventor or co-inventor on 19 patent applications spanning automotive chassis, steering, suspension, vehicle-state estimation, trailer assistance, and mechanical systems.",
    ]
    for item in publications:
        add_bullet(doc, bullet_num_id, item, small=True)

    doc.core_properties.title = "Kyungtack Lee - Resume"
    doc.core_properties.subject = "Vehicle Motion Control and Autonomous Driving"
    doc.core_properties.author = "Kyungtack Lee"
    doc.core_properties.keywords = (
        "motion planning, vehicle control, autonomous driving, real-time, NVIDIA"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_docx()
